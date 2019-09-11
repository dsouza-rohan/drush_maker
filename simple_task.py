import os
import sys
import subprocess
import acapi
from docx import Document
import yaml
import datetime
import time


class AcqUtility:
    """Simple AcqUtility utility class"""

    action = ""  # action to take
    acq_cli = 0
    acq_sub = ""
    acq_env = ""
    acq_site = 0
    allowed = ["dev", "test", "dev6"]
    commands = ["coder", "otl", "otl_doc", "mtp_doc", "db_dump", "yaml_valid", "debug_ssh"]
    pub_key_path = ""
    pub_key_pass = ""
    local_dir = ""
    mode = "drush"  # drush | ssh
    run_setting = ""

    def __init__(self, action, sub, env):
        self.action = action
        self.acq_env = env
        self.acq_sub = sub

        if action in self.commands :
            self.run_setting = self.get_run_settings()
            if self.run_setting == "":
                print("To run the script create a simple_task.yaml file")
            else:
                if "my_key_path" in self.run_setting.keys():
                    self.pub_key_path = self.run_setting.get('my_key_path')
                else:
                    sys.exit("my_key_path is not set in yaml")

                if "allowed_env" in self.run_setting.keys():
                    self.allowed = self.run_setting.get('allowed_env')
                else:
                    sys.exit("allowed_env is not set in yaml")
        else:
            sys.exit("Invalid action; {act}".format(act=action))

        print("start--")

    def __del__(self):
        print("done--")

    def mk_acq_client(self):

        if self.acq_env in self.allowed:

            drush_alias = "@{sub}.{env}".format(sub=self.acq_sub, env=self.acq_env)

            if self.mode == "drush":
                return drush_alias

            elif self.mode == "ssh":
                sys.exit("No ssh mode available...yet")
        else:
            sys.exit("Connection to {} is not allowed".format(self.acq_env))
            # ==>Available envs are !!
            # for env in envs:
            #     print "Env: {env} SSH Host: {host}".format(env=env, host=envs[env]['ssh_host'])

    def get_acq_site_name(self):
        c = acapi.Client()
        site = c.site(self.acq_sub)
        domains = site.environment(self.acq_env).domains()
        domain_str = ""
        for domain in domains:
            domain_str += "{},".format(domain)

        return domain_str.strip(",")

    def get_run_settings(self):
        setting = ""
        if os.path.exists('simple_task.yaml'):
            with open("simple_task.yaml", 'r') as stream:
                try:
                    setting = yaml.safe_load(stream)
                except yaml.YAMLError as exc:
                    print(exc)

        return setting

    def run_scripts(self, action):

        if action == "coder":
            self.mk_coder_report()

        elif action == "otl":
            self.mk_site_uli()

        elif action == "otl_doc":
            self.mk_site_uli_doc()

        elif action == "mtp_doc":
            print(action)

        elif action == "db_dump":
            print(action)
            # todo: db dump...

        elif action == "yaml_valid":
            commands = self.get_run_settings()
            print(commands)

        elif action == "debug_ssh":
            print(action)
            self.run_ssh_agent()

        # todo:  and 5x b-fix queries...?

    def mk_coder_report(self):
        alias = self.mk_acq_client()
        commands = self.run_setting

        if "coder_report" in commands.keys():
            coder_report = commands.get('coder_report')
            print(coder_report)

            for todo_text, cmd in coder_report.items():
                print("{todo}-->{cmd}".format(todo=todo_text, cmd=cmd.format(alias=alias)))
                process = subprocess.Popen([cmd.format(alias=alias)], stdout=subprocess.PIPE, shell=True)
                (out, err) = process.communicate()
                print(out)
                # todo: save coder output to file

    def mk_site_uli(self):

        alias = self.mk_acq_client()

        os.system("drush {alias} uli 1 --no-browser".format(alias=alias))

    def mk_site_uli_doc(self):

        alias = self.mk_acq_client()
        site_names = self.get_acq_site_name()
        commands = self.run_setting
        now_time = datetime.datetime.fromtimestamp(time.time()).strftime('%Y-%m-%d_%H:%M:%S')
        save_file = 'otl_doc_{sub}_{date}.docx'.format(sub=self.acq_sub, date=now_time)

        if "one_time_link" in commands.keys():
            commands = commands.get('one_time_link')

            document = Document()
            document.add_heading('OTL Document -- {}'.format(self.acq_sub), 0)
            document.add_paragraph('Site Name: {name}'.format(name=site_names))
            document.add_paragraph('Subscription: {sub}'.format(sub=self.acq_sub))

            for todo_text, cmd in commands.items():
                print(cmd.format(alias=alias))
                process = subprocess.Popen([cmd.format(alias=alias)], stdout=subprocess.PIPE, shell=True)
                (out, err) = process.communicate()
                print(out)

                document.add_paragraph(todo_text)
                font = document.add_paragraph("output:\n  {}".format(out))
                font.highlight_color = "red"
                document.add_page_break()

            document.save(save_file)

            # todo: check ssl certificate of prod site.

        else:
            print("Config 'one_time_link' missing, simple_task.yaml file")
            return False

    def run_ssh_agent(self):
        os.system("eval $(ssh-agent -s)")
        os.system("ssh-add {path}".format(path=self.pub_key_path))


def runner_handler(event):
    if len(event) > 0:
        bf = AcqUtility
        #  pass target folder/repo to script
        bf.run_scripts(AcqUtility(event['action'], event['sub'], event['env']), event['action'])


print(sys.argv)

if len(sys.argv) == 4:
    event_local = {'action':  sys.argv[1], 'sub':  sys.argv[2], 'env':  sys.argv[3]}
    runner_handler(event_local)
else:
    if len(sys.argv) == 2:
        debug = ["debug_ssh", "yaml_valid"]
        if sys.argv[1] == "--help":

            print("\n  simple_task.py [action] [subscription] [environment] \n" \
                  " action--> [coder| otl| otl_doc| db_dump ..]\n" \
                  " subscription--> Subscription of the site\n" \
                  " environment--> [dev | test]\n " \
                  "\nOR debug_ssh to solve ssh-agent issue \n" \
                  "yaml_valid , to validate yaml file")

        elif sys.argv[1] in debug:

            event_local = {'action': sys.argv[1], 'sub': 0, 'env': 0}
            runner_handler(event_local)

        else:
            print("3 parameters required. Please use --help")
    else:
        print("use --help")

    # Local testing
    event_local = {'action': "", 'sub': "", 'env':  "dev"}
    runner_handler(event_local)

